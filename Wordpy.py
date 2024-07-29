from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

def Titulo_Terminal(a):
    print('=' * 30)
    print(a)
    print('=' * 30)
    return a

def linha():
    print('-' * 30)

def mostrar_exp(mostrar):
    if len(mostrar) == 0:
        return ' '
    else:
        return '\n'.join(mostrar)

doc = Document()
Titulo_Terminal('   Crie seu curriculo!')
name = input('Escreva o seu nome completo: ')
idade = input('Escreva sua idade: ')
tel = input('Escreva seu telefone/celular: ')
email = input('Escreva seu email: ')
escolaridade = []
experiencias = []
idiomas = []
linha()
f1 = int(input('Quantas formações academicas gostaria de adicionar: '))
linha()
for academic in range(f1):
    formacao = input(f'Me diga sua {academic + 1}ª formação: ')
    escolaridade.append(formacao)
linha()
f2 = int(input('Quantas experiencias anteriores possui: '))
linha()
for exp in range(f2):
    e1 = input(f'Me diga sua {exp + 1}ª experiencia: ')
    experiencias.append(e1)
linha()
f3 = int(input('Quantos idiomas sabe falar (se for apenas o português, digite 0): '))
linha()
if f3 > 0:
    for idi in range(f3):
        e2 = input(f'Me diga seu {idi + 1}º idioma: ')
        idiomas.append(e2)
linha()
print('Muito Obrigado! Seu curriculo vai ser criado!')

# Dados
data = {
    'Titulo': f'{name}',
    'Paragrafo': 'Nacionalidade: Brasileiro(a)',
    'Idade': f'Idade: {idade}',
    'Email': f'Email: {email}',
    'Telefone': f'Telefone: {tel}',
    'Titulo2': 'Formação Acadêmica',
    'Escolaridade': mostrar_exp(escolaridade),
    'Titulo3': 'Experiência Profissional',
    'Experiencias': mostrar_exp(experiencias),
    'Titulo4': 'Idiomas',
    'Idiomas': mostrar_exp(idiomas),
    'Titulo5': 'Informações Complementares',
    'Informações': 'Proatividade, Disciplina e comunicação empatica'
}

# Estilos
styles = doc.styles

# Estilo de parágrafo
p_style = styles.add_style('Paragraph', WD_STYLE_TYPE.PARAGRAPH)
p_style.font.name = 'Times New Roman'
p_style.font.size = Pt(11)
p_style.font.bold = False

# Estilo do título principal
head_style = styles.add_style('Head', WD_STYLE_TYPE.PARAGRAPH)
head_style.font.name = 'Times New Roman'
head_style.font.size = Pt(24)
head_style.font.color.rgb = RGBColor(0, 0, 0)
head_style.font.bold = True
head_style.font.underline = True

# Estilo dos subtítulos
subhead_style = styles.add_style('SubHead', WD_STYLE_TYPE.PARAGRAPH)
subhead_style.font.name = 'Times New Roman'
subhead_style.font.size = Pt(16)
subhead_style.font.bold = True
subhead_style.font.color.rgb = RGBColor(0, 0, 0)

# Adicionar conteúdo ao documento com os estilos aplicados
doc.add_paragraph(data['Titulo'], style='Head')
doc.add_paragraph(data['Paragrafo'], style='Paragraph')
doc.add_paragraph(data['Idade'], style='Paragraph')
doc.add_paragraph(data['Email'], style='Paragraph')
doc.add_paragraph(data['Telefone'], style='Paragraph')

doc.add_paragraph(data['Titulo2'], style='SubHead')
doc.add_paragraph(data['Escolaridade'], style='Paragraph')

doc.add_paragraph(data['Titulo3'], style='SubHead')
doc.add_paragraph(data['Experiencias'], style='Paragraph')

doc.add_paragraph(data['Titulo4'], style='SubHead')
doc.add_paragraph(data['Idiomas'], style='Paragraph')

doc.add_paragraph(data['Titulo5'], style='SubHead')
doc.add_paragraph(data['Informações'], style='Paragraph')

# Salvar o documento
doc.save(f'Curriculo_{name}.docx')





